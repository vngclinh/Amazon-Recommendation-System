# -*- coding: utf-8 -*-
"""
Insert available charts into the course report (docx), mark missing charts with a red X,
and set the whole document to Times New Roman 13pt.
Run: .venv/Scripts/python.exe docs/insert_figures.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOC = os.path.join(ROOT, "docs", "Untitled document.docx")

# fignum -> (image path or None, title). None => red-X placeholder.
FIGS = {
    "3.1": (r"goodreads/charts/pipeline-data-processing.jpg", "Overall data processing pipeline"),
    "3.2": (r"goodreads/charts/newplot2.png",                 "Goodreads rating distribution"),
    "3.3": (r"goodreads/charts/newplot3.png",                 "Review volume over time"),
    "3.4": (r"goodreads/charts/newplot4.png",                 "Review length by rating"),
    "3.5": (r"goodreads/charts/newplot6.png",                 "Book and author popularity distribution"),
    "4.1": (None,                                             "Overall recommendation pipeline"),
    "4.2": (r"src/visualize_data/amazon_recsys_pipeline.png", "Amazon five-stage pipeline"),
    "4.3": (r"goodreads/goodreads_recsys_pipeline.png",       "Goodreads four-phase pipeline"),
    "5.1": (None, "Comparison of Amazon rating-prediction models"),
    "5.2": (None, "Amazon full-catalog ranking failure"),
    "5.3": (None, "Edge weighting ablation results on Goodreads"),
    "5.4": (None, "sampled@500 and full-ranking comparison"),
    "5.5": (None, "Web demo interface"),
}

FONT = "Times New Roman"
BODY_PT = 13
IMG_W = Inches(6.0)
RED = RGBColor(0xC0, 0x00, 0x00)


def set_run_font(run, size=BODY_PT, name=FONT, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(a), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def new_par_after(par):
    p = OxmlElement("w:p")
    par._p.addnext(p)
    return Paragraph(p, par._parent)


def new_par_before(par):
    p = OxmlElement("w:p")
    par._p.addprevious(p)
    return Paragraph(p, par._parent)


def fill_image(par, fignum):
    """Turn `par` into a centered image (or red-X) paragraph for `fignum`."""
    path, title = FIGS[fignum]
    par.clear()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path:
        run = par.add_run()
        run.add_picture(os.path.join(ROOT, path.replace("/", os.sep)), width=IMG_W)
    else:
        r = par.add_run("✗")          # heavy ballot X
        set_run_font(r, size=48, color=RED)
        r.bold = True
    return par


def fill_caption(par, fignum):
    path, title = FIGS[fignum]
    par.clear()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lab = par.add_run(f"Figure {fignum}. ")
    set_run_font(lab, size=BODY_PT)
    lab.bold = True
    cap = par.add_run(title)
    set_run_font(cap, size=BODY_PT)
    cap.italic = True
    if path is None:
        note = par.add_run("   [ cần thêm hình — insert chart here ]")
        set_run_font(note, size=BODY_PT, color=RED)
        note.bold = True
    return par


def place_at_anchor(anchor, fignum):
    """Reuse `anchor` as the image paragraph, add caption after. Returns caption par."""
    fill_image(anchor, fignum)
    cap = new_par_after(anchor)
    fill_caption(cap, fignum)
    return cap


def place_after(anchor, fignum):
    img = new_par_after(anchor)
    fill_image(img, fignum)
    cap = new_par_after(img)
    fill_caption(cap, fignum)
    return cap


def place_before(target, fignum):
    img = new_par_before(target)
    fill_image(img, fignum)
    cap = new_par_after(img)
    fill_caption(cap, fignum)
    return cap


def find_par(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    return None


doc = Document(DOC)

# ---------- 1) Replace [Insert Figure ...] placeholders ----------
for p in list(doc.paragraphs):
    t = p.text
    if "[Insert Figure 3.1" in t:
        place_at_anchor(p, "3.1")
    elif "[Insert Figure 3.2" in t:                       # combined 3.2/3.3/3.4/3.5
        anchor = place_at_anchor(p, "3.2")
        anchor = place_after(anchor, "3.3")
        anchor = place_after(anchor, "3.4")
        anchor = place_after(anchor, "3.5")
    elif "[Insert Figure 4.1" in t:
        place_at_anchor(p, "4.1")
    elif "[Insert Figure 5.5" in t:
        place_at_anchor(p, "5.5")
    elif "[Insert Figure 5.2" in t:
        place_at_anchor(p, "5.2")
    elif "[Insert Figure 5.3" in t:
        place_at_anchor(p, "5.3")
    elif "[Insert Figure 5.4" in t:
        place_at_anchor(p, "5.4")

# ---------- 2) Insert figures that had no placeholder ----------
# Fig 4.2 (Amazon pipeline) -> before heading "4.3. Phase 2"
h43 = find_par(doc, lambda s: s.strip().startswith("4.3. Phase 2"))
if h43 is not None:
    place_before(h43, "4.2")

# Fig 4.3 (Goodreads pipeline) -> before heading "4.4. Edge Weighting"
h44 = find_par(doc, lambda s: s.strip().startswith("4.4. Edge Weighting"))
if h44 is not None:
    place_before(h44, "4.3")

# Fig 5.1 (Amazon model comparison) -> before heading "5.3. Limitations"
h53 = find_par(doc, lambda s: s.strip().startswith("5.3. Limitations"))
if h53 is not None:
    place_before(h53, "5.1")

# ---------- 3) Global font: Times New Roman, body 13pt ----------
# Normal style default
try:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
except KeyError:
    pass


def is_heading(par):
    try:
        n = (par.style.name or "").lower()
    except Exception:
        n = ""
    return n.startswith("heading") or n.startswith("title")


def apply_paragraph(par):
    head = is_heading(par)
    for run in par.runs:
        # keep heading sizes (visual hierarchy) but force TNR; body -> 13
        if head:
            set_run_font(run, size=None)
        else:
            # don't shrink our big red X
            cur = run.font.size
            if cur is not None and cur > Pt(20):
                set_run_font(run, size=None)
            else:
                set_run_font(run, size=BODY_PT)


for par in doc.paragraphs:
    apply_paragraph(par)

for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    set_run_font(run, size=BODY_PT)

doc.save(DOC)
print("OK -> saved", DOC)
